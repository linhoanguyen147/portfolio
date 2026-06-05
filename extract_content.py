"""
Script để trích xuất nội dung từ các file NV1-NV6
Chạy: python extract_content.py
"""
import os
import json

BASE = r"C:\Users\ADMIN\.gemini\antigravity\scratch\portfolio"
results = {}

# --- Đọc file DOCX ---
try:
    import docx
    for fname in ["nv1.docx", "nv2.docx", "nv3.docx"]:
        fpath = os.path.join(BASE, fname)
        if os.path.exists(fpath):
            doc = docx.Document(fpath)
            lines = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    # Lấy style để biết heading hay normal
                    style = para.style.name if para.style else "Normal"
                    lines.append({"style": style, "text": t})
            # Đọc tables
            tables = []
            for table in doc.tables:
                tbl = []
                for row in table.rows:
                    tbl.append([cell.text.strip() for cell in row.cells])
                tables.append(tbl)
            results[fname] = {"paragraphs": lines, "tables": tables}
            print(f"[OK] {fname}: {len(lines)} paragraphs, {len(tables)} tables")
except ImportError:
    print("[ERROR] python-docx not installed")

# --- Đọc file PDF ---
try:
    import PyPDF2
    for fname in ["nv4.pdf", "nv5.pdf", "nv6.pdf"]:
        fpath = os.path.join(BASE, fname)
        if os.path.exists(fpath):
            with open(fpath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append({"page": i+1, "text": text.strip()})
            results[fname] = {"pages": pages}
            total_text = sum(len(p["text"]) for p in pages)
            print(f"[OK] {fname}: {len(pages)} pages, {total_text} chars")
except ImportError:
    print("[ERROR] PyPDF2 not installed")

# --- Lưu kết quả ---
out = os.path.join(BASE, "extracted_content.json")
with open(out, "w", encoding="utf-8") as f:
    json.dump(results, f, ensure_ascii=False, indent=2)
print(f"\n[DONE] Saved to extracted_content.json")

# --- In nhanh nội dung để xem ---
for fname, content in results.items():
    print(f"\n{'='*60}")
    print(f"FILE: {fname}")
    print('='*60)
    if "paragraphs" in content:
        for item in content["paragraphs"][:50]:
            prefix = "## " if "Heading" in item["style"] else "   "
            print(f"{prefix}{item['text']}")
    elif "pages" in content:
        for pg in content["pages"][:3]:
            print(f"\n--- Page {pg['page']} ---")
            print(pg["text"][:2000])
